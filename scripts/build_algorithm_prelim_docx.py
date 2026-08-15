from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


TEMPLATE = Path("docs/templates/AI_for_research_algorithm_preliminary_template.docx")
OUTPUT = Path("outputs/documents/GOAI-VCell-Route_算法赛初赛方案.docx")

NAVY = "17375E"
BLUE = "2E74B5"
TEAL = "168C88"
ORANGE = "D97832"
INK = "243447"
MUTED = "5E6B78"
PALE_BLUE = "EAF2F8"
PALE_TEAL = "E8F5F3"
PALE_ORANGE = "FFF1E7"
PALE_GRAY = "F3F5F7"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges: dict[str, str]) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, options in edges.items():
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in options.items():
            element.set(qn("w:" + key), str(value))


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run_font(run, name="微软雅黑", size=10.5, bold=None, color=INK) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_text(paragraph, text: str, *, bold=False, color=INK, size=10.5, name="微软雅黑"):
    run = paragraph.add_run(text)
    set_run_font(run, name=name, size=size, bold=bold, color=color)
    return run


def set_para(
    paragraph,
    *,
    before=0,
    after=4,
    line=1.25,
    alignment=None,
    keep_with_next=False,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if alignment is not None:
        paragraph.alignment = alignment
    fmt.keep_with_next = keep_with_next


def add_body(doc, text: str, *, bold_prefix: str | None = None, after=5):
    p = doc.add_paragraph()
    set_para(p, after=after, line=1.34)
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, bold=True, color=NAVY)
        add_text(p, text[len(bold_prefix) :])
    else:
        add_text(p, text)
    return p


def add_bullet(doc, text: str, *, level=0, color=INK):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    set_para(p, after=2, line=1.25)
    add_text(p, text, color=color)
    return p


def add_number(doc, number: int, title: str, body: str):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(0.9)
    table.columns[1].width = Cm(15.9)
    left, right = table.rows[0].cells
    set_cell_shading(left, TEAL)
    set_cell_shading(right, PALE_TEAL)
    for cell in (left, right):
        set_cell_margins(cell, top=95, bottom=95)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_border(
            cell,
            top={"val": "single", "sz": "6", "color": "D0E6E3"},
            bottom={"val": "single", "sz": "6", "color": "D0E6E3"},
            start={"val": "single", "sz": "6", "color": "D0E6E3"},
            end={"val": "single", "sz": "6", "color": "D0E6E3"},
        )
    lp = left.paragraphs[0]
    set_para(lp, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_text(lp, str(number), bold=True, color=WHITE, size=12)
    rp = right.paragraphs[0]
    set_para(rp, after=0, line=1.2)
    add_text(rp, title + "  ", bold=True, color=NAVY)
    add_text(rp, body, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_heading(doc, text: str, level: int):
    p = doc.add_heading(text, level=level)
    set_para(
        p,
        before=12 if level == 1 else 7,
        after=5,
        line=1.1,
        keep_with_next=True,
    )
    for run in p.runs:
        set_run_font(
            run,
            size=16 if level == 1 else 12.5,
            bold=True,
            color=BLUE if level == 1 else NAVY,
        )
    if level == 1:
        p.paragraph_format.page_break_before = True
    return p


def add_callout(doc, title: str, body: str, *, fill=PALE_ORANGE, accent=ORANGE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, start=170, bottom=150, end=170)
    set_cell_border(
        cell,
        start={"val": "single", "sz": "18", "color": accent},
        top={"val": "single", "sz": "3", "color": fill},
        bottom={"val": "single", "sz": "3", "color": fill},
        end={"val": "single", "sz": "3", "color": fill},
    )
    p = cell.paragraphs[0]
    set_para(p, after=0, line=1.28)
    add_text(p, title + "  ", bold=True, color=accent)
    add_text(p, body)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_table(doc, headers, rows, widths=None, alignments=None, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    if widths:
        for idx, width in enumerate(widths):
            table.columns[idx].width = Cm(width)
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    for idx, value in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, top=105, bottom=105)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        set_para(p, after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(p, str(value), bold=True, color=WHITE, size=font_size)
    for row_idx, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            set_cell_shading(cell, WHITE if row_idx % 2 == 0 else PALE_GRAY)
            set_cell_margins(cell, top=88, bottom=88)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            alignment = WD_ALIGN_PARAGRAPH.LEFT
            if alignments and idx < len(alignments):
                alignment = alignments[idx]
            set_para(p, after=0, line=1.18, alignment=alignment)
            add_text(p, str(value), size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_architecture(doc):
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [3.5, 4.5, 4.4, 4.4]
    for idx, width in enumerate(widths):
        table.columns[idx].width = Cm(width)

    content = [
        (
            ("输入", NAVY, WHITE),
            ("细胞与培养条件\n菌株 / 培养基 / 温度 / 时间", PALE_BLUE, INK),
            ("观测上下文\n来源 / 仪器 / 板号", PALE_GRAY, INK),
            ("化合物结构\nMorgan-2048", PALE_TEAL, INK),
        ),
        (
            ("表示学习", NAVY, WHITE),
            ("共享 M6 编码器", PALE_BLUE, INK),
            ("独立 Calibration 编码器", PALE_GRAY, INK),
            ("冻结 OP3 RNA 编码器\n+ context gate", PALE_TEAL, INK),
        ),
        (
            ("三项分解", NAVY, WHITE),
            ("背景状态 B₆\n+ 基础响应 R₆", PALE_BLUE, INK),
            ("测量校准 C₆", PALE_GRAY, INK),
            ("迁移响应 R₉.₆", PALE_TEAL, INK),
        ),
        (
            ("响应融合", NAVY, WHITE),
            ("R₁₂ = 1.075R₉.₆ - 0.075R₆\n+ 0.15·I(|R₆|≥0.5)·(R₆-blend)", PALE_ORANGE, INK),
            ("仅在 R10 / S1 启用", PALE_ORANGE, INK),
            ("大效应位置向 R₆ 小幅回拉", PALE_ORANGE, INK),
        ),
        (
            ("输出", NAVY, WHITE),
            ("ŷ = B₆ + C₆ + R₁₂", BLUE, WHITE),
            ("4,422 维完整 log₂ 蛋白丰度", BLUE, WHITE),
            ("逐行 support router", BLUE, WHITE),
        ),
    ]
    for r_idx, row_values in enumerate(content):
        row = table.rows[r_idx]
        prevent_row_split(row)
        for c_idx, (value, fill, color) in enumerate(row_values):
            cell = row.cells[c_idx]
            set_cell_shading(cell, fill)
            set_cell_margins(cell, top=150, bottom=150, start=120, end=120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(
                cell,
                top={"val": "single", "sz": "6", "color": WHITE},
                bottom={"val": "single", "sz": "6", "color": WHITE},
                start={"val": "single", "sz": "6", "color": WHITE},
                end={"val": "single", "sz": "6", "color": WHITE},
            )
            p = cell.paragraphs[0]
            set_para(p, after=0, line=1.15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            add_text(p, value, bold=c_idx == 0 or r_idx == 4, color=color, size=9.2)
    caption = doc.add_paragraph()
    set_para(caption, before=3, after=5, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(caption, "图 1  GOAI-M12.0 分解式架构（阶段性候选）", color=MUTED, size=9)


def add_page_number(paragraph):
    add_text(paragraph, "GOAI AI for Research赛道 | 算法赛初赛方案  ·  ", color=MUTED, size=8.5)
    add_text(paragraph, "第 ", color=MUTED, size=8.5)
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)
    set_run_font(run, size=8.5, color=MUTED)
    add_text(paragraph, " 页", color=MUTED, size=8.5)


def clear_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(4)

    for style_name in ("List Bullet", "List Bullet 2"):
        style = doc.styles[style_name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(10.2)
        style.font.color.rgb = RGBColor.from_string(INK)

    for style_name, size, color in (
        ("Heading 1", 16, BLUE),
        ("Heading 2", 12.5, NAVY),
    ):
        style = doc.styles[style_name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def build_document() -> Document:
    doc = Document(TEMPLATE)
    clear_body(doc)
    configure_styles(doc)

    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(2.05)
    section.right_margin = Cm(2.05)
    section.header_distance = Cm(0.75)
    section.footer_distance = Cm(0.7)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    set_para(fp, after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_page_number(fp)

    props = doc.core_properties
    props.title = "GOAI-VCell Route：面向未见实体的分解式酵母扰动响应预测"
    props.subject = "GOAI AI for Research 算法赛初赛方案"
    props.keywords = "GOAI, virtual cell, yeast, perturbation, proteomics, OOD"

    # Cover
    p = doc.add_paragraph()
    set_para(p, before=18, after=4, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "GOAI · AI for Research", bold=True, color=TEAL, size=12)

    p = doc.add_paragraph()
    set_para(p, before=25, after=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "GOAI-VCell Route", bold=True, color=NAVY, size=27)

    p = doc.add_paragraph()
    set_para(p, after=20, alignment=WD_ALIGN_PARAGRAPH.CENTER, line=1.2)
    add_text(p, "面向未见实体的分解式酵母扰动响应预测", bold=True, color=BLUE, size=16)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(4.1)
    table.columns[1].width = Cm(10.8)
    cover_rows = [
        ("参赛类型", "算法赛题"),
        ("参赛方向", "方向一：虚拟细胞（开放知识榜）"),
        ("阶段模型", "GOAI-M12.0 · 本地严格 OOF 阶段性候选"),
        ("团队信息", "【团队名称 / 成员信息待团队填写】"),
    ]
    for idx, (label, value) in enumerate(cover_rows):
        for col, text in enumerate((label, value)):
            cell = table.rows[idx].cells[col]
            set_cell_shading(cell, PALE_BLUE if col == 0 else WHITE)
            set_cell_margins(cell, top=140, bottom=140)
            set_cell_border(
                cell,
                top={"val": "single", "sz": "5", "color": "D8E2EC"},
                bottom={"val": "single", "sz": "5", "color": "D8E2EC"},
                start={"val": "single", "sz": "5", "color": "D8E2EC"},
                end={"val": "single", "sz": "5", "color": "D8E2EC"},
            )
            p = cell.paragraphs[0]
            set_para(p, after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER if col == 0 else WD_ALIGN_PARAGRAPH.LEFT)
            add_text(p, text, bold=col == 0, color=NAVY if col == 0 else INK, size=10.5)

    p = doc.add_paragraph()
    set_para(p, before=22, after=3, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "初赛方案说明文档 · 内部复核稿", bold=True, color=ORANGE, size=11)
    p = doc.add_paragraph()
    set_para(p, after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "版本日期：2026 年 8 月 15 日", color=MUTED, size=9.5)

    doc.add_page_break()

    # Executive summary
    p = doc.add_paragraph()
    set_para(p, before=2, after=6)
    add_text(p, "方案一句话", bold=True, color=TEAL, size=12)
    add_callout(
        doc,
        "核心判断",
        "完整蛋白质组预测不是一个单块回归问题。我们把预测拆成细胞背景、观测校准与扰动响应，"
        "再根据菌株和化合物在训练中的支持情况逐行路由；只有证据充分的模块才进入对应场景。",
        fill=PALE_TEAL,
        accent=TEAL,
    )

    add_number(doc, 1, "先拆开",
               "用 B（背景状态）+ C（测量校准）+ R（扰动响应）解释绝对蛋白丰度，避免批次信息与生物效应互相代偿。")
    add_number(doc, 2, "再路由",
               "已见菌株 / 未见化合物、未见菌株 / 已见化合物、双未知与时间外推使用不同的可靠路径。")
    add_number(doc, 3, "只保留有证据的迁移",
               "OP3 RNA 扰动表示仅用于验证有效的新化合物场景；菌株语义、已见菌株专家和双未知联合语义未过门禁，不强行并入。")

    add_table(
        doc,
        ["阶段结论", "当前证据"],
        [
            ("本地最优候选", "GOAI-M12.0；FC PCC 0.426342（严格 OOF proxy，非官方分数）"),
            ("相对上一版", "FC +0.000203，37 个 chemical cluster 配对 bootstrap 95% CI 为 [+0.000059, +0.000355]"),
            ("模型边界", "FC 与高效应 PCC 上升，但 Context PCC 下降；它是 FC 优先候选，不是全指标支配模型"),
            ("正式提交前", "按最新允许标签边界重训；确认 4,422 / 5,243 输出列契约；完整披露 OP3 外部数据"),
        ],
        widths=[4.1, 12.4],
        font_size=9.5,
    )

    # Section 1
    add_heading(doc, "一、项目概述", 1)
    add_heading(doc, "1.1 项目名称", 2)
    add_body(doc, "GOAI-VCell Route：面向未见实体的分解式酵母扰动响应预测。")

    add_heading(doc, "1.2 参赛方向", 2)
    add_body(doc, "算法赛题，方向一“虚拟细胞（虚拟酵母扰动响应预测）”。由于方案使用公开 RNA 扰动数据进行预训练，本项目按开放知识榜准备与披露。")

    add_heading(doc, "1.3 方案概述", 2)
    add_body(
        doc,
        "任务要求根据菌株、化合物、培养条件和测量上下文，预测测试样本的完整 log₂ 蛋白丰度向量。真正困难的不是在已知组合上拟合均值，而是在未见化合物、未见菌株及二者同时未知时，仍给出方向正确、幅度稳定的扰动响应。",
    )
    add_body(
        doc,
        "我们的方案采用“分解 + 路由 + 证据门禁”。分解负责区分细胞本底、测量偏移和扰动效应；路由让不同泛化场景使用相应模型；证据门禁要求每个新增模块通过同协议 OOF、配对 bootstrap 和负对照后才能进入最终路径。当前阶段性候选 GOAI-M12.0 只替换证据最充分的 S1 路由，其余场景保留更稳健的冻结模型。",
    )

    # Section 2
    add_heading(doc, "二、科学问题理解", 1)
    add_heading(doc, "2.1 科学问题与研究对象", 2)
    add_body(
        doc,
        "研究对象是酿酒酵母在遗传背景、化学扰动、培养基、温度和作用时间共同影响下的蛋白质组状态。输出不是一个标签，而是数千个相互关联的蛋白丰度；评分既考察绝对谱形，也考察相对匹配对照的 fold change、高效应蛋白和实体外推。",
    )
    add_body(doc, "我们将难点归纳为三类：")
    add_bullet(doc, "高维且稀疏：蛋白标签存在缺失，强响应蛋白占比低，普通均方误差容易被背景丰度主导。")
    add_bullet(doc, "实体分布外推：新化合物和新菌株没有可直接学习的 ID 参数，必须依赖可迁移结构或稳定回退。")
    add_bullet(doc, "生物效应与测量偏移共存：仪器、来源和板号能解释观测差异，但不应替代化合物响应。")

    add_heading(doc, "2.2 科学意义", 2)
    add_body(
        doc,
        "一个可信的虚拟细胞模型应回答“在这个菌株和环境中，某个新扰动会把细胞推向什么蛋白状态”，而不只是复述训练集平均值。若模型能稳定筛出方向正确的高效应蛋白，可用于缩小湿实验候选范围、辅助比较菌株改造方案，并为后续作用机制分析提供可检验的蛋白响应假设。",
    )
    add_callout(
        doc,
        "研究边界",
        "本项目首先解决预测与泛化，不把相关性输出直接解释为因果机制。机制结论需要独立实验、通路证据或干预验证。",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )

    # Section 3
    add_heading(doc, "三、技术方案与预期方法路线", 1)
    add_heading(doc, "3.1 技术方案", 2)
    add_body(
        doc,
        "核心输出被写成 ŷ = B + C + R。B 表示由菌株与培养条件决定的背景蛋白状态；C 只读取来源、仪器与板号，用于吸收观测偏移；R 表示相对背景的化学扰动响应。这样做让三个分支各自承担明确职责，也便于在 OOD 场景中替换单一部件。",
    )
    add_architecture(doc)

    add_body(
        doc,
        "M6 主干用共享的 cell/context 编码器同时产生背景 B₆ 与基础响应 R₆。独立校准分支产生 C₆。M9.6 将 Morgan-2048 化合物指纹输入由 OP3 单细胞扰动数据预训练的冻结 RNA 编码器，再经上下文门控映射为 4,422 维蛋白响应 R₉.₆。",
    )
    add_body(
        doc,
        "M12.0 是一个固定高效应 specialist：先用 M9.6 为主进行响应融合，再在 |R₆|≥0.5 的位置向 M6 小幅回拉。阈值只读取模型预测，不读取验证标签；融合参数由训练阶段 OOF 冻结，测试阶段不再调整。",
    )

    add_heading(doc, "3.2 逐场景路由", 2)
    add_table(
        doc,
        ["路由", "含义", "测试行数", "采用路径", "选择理由"],
        [
            ("R10 / S1", "菌株已见、化合物未知", "2,072", "GOAI-M12.0", "OP3 化合物迁移在严格 chemical-held-out OOF 中有效"),
            ("R01 / S2", "菌株未知、化合物已见", "1,594", "M5.2 / M2", "SNP-MDS 有信号，但加入主模型后 FC 下降"),
            ("R00 / S3", "菌株与化合物均未知", "425", "M5.2 / M2", "双语义模型未超过 zero-semantic 与冻结基线"),
            ("R11", "菌株与化合物均已见", "135", "M5.2 time route", "保留稳定的时间外推路径"),
            ("Control / QC", "对照与质量控制", "228", "Background / Control", "不注入处理响应"),
        ],
        widths=[2.0, 3.7, 1.6, 3.1, 6.1],
        alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
        font_size=8.6,
    )
    add_body(
        doc,
        "逐行路由不使用官方 split 名称作为捷径，而是依据当前 checkpoint 的 strain / chemical support vocabulary 判断实体是否已见。五类行数合计 4,454，与测试样本总数一致。",
    )

    add_heading(doc, "3.3 训练、验证与推理流程", 2)
    for number, title, body in (
        (1, "数据合同", "按 sample_ID 连接元数据与蛋白矩阵；仅在允许的训练标签中拟合缺失过滤、均值、尺度和实体词表。"),
        (2, "严格 OOF", "按完整化合物簇留出构造 S1 folds，使用 seeds 42 / 43 / 2026；模型、融合权重和阈值在相同 fold 合同下比较。"),
        (3, "负对照门禁", "真实外部表示必须优于 shuffled / zero；新增分支还必须在与父模型融合后改善目标指标。"),
        (4, "冻结与推理", "冻结参数和路由，对 4,454 个 test 样本生成绝对 log₂ 蛋白丰度，并校验 ID、列序、有限值和哈希。"),
    ):
        add_number(doc, number, title, body)

    add_heading(doc, "3.4 数据来源、依赖工具与运行环境", 2)
    add_table(
        doc,
        ["类别", "来源与用途", "许可 / 边界"],
        [
            ("官方数据", "GOAI train_val / test：元数据、测量上下文与蛋白标签", "非公开赛事数据；仅限竞赛范围使用，不随代码发布"),
            ("外部数据", "Open Problems Single-Cell Perturbations：2023-09-12_de_by_cell_type_train.h5ad，用于 RNA 扰动编码器预训练", "CC BY 4.0；文件 SHA256 以 manifest 记录"),
            ("化学表征", "RDKit Morgan fingerprint，radius=2，2,048 bits", "开源依赖；化合物映射与缺失回退写入审计表"),
            ("训练框架", "Python、PyTorch、NumPy、pandas、scikit-learn", "核心流程不依赖商业 API 或闭源模型"),
        ],
        widths=[2.4, 9.1, 5.0],
        font_size=8.9,
    )

    # Section 4
    add_heading(doc, "四、阶段性实验结果或可行性验证", 1)
    add_heading(doc, "4.1 验证设计", 2)
    add_body(
        doc,
        "阶段性比较采用固定的 S1 / R10 chemical-held-out OOF：5,078 个 treatment 样本，37 个完整 chemical clusters，三个随机种子等权。所有数值均为本地 proxy，不是官方 PSS，也不与不同验证协议的结果混比。",
    )

    add_heading(doc, "4.2 当前结果", 2)
    add_table(
        doc,
        ["模型", "FC PCC", "Context PCC", "High PCC", "High F1", "Abs R²"],
        [
            ("M6 core", "0.371973", "0.098530", "0.635746", "0.182687", "0.979300"),
            ("M11 blend", "0.426139", "0.062244", "0.600870", "0.233938", "0.979193"),
            ("M12.0 specialist", "0.426342", "0.060967", "0.603184", "0.233970", "0.979150"),
        ],
        widths=[4.0, 2.4, 2.7, 2.7, 2.5, 2.5],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 5,
        font_size=9.1,
    )
    add_body(
        doc,
        "相对 M11，M12.0 的 FC PCC 提升 +0.000203；按 37 个 chemical clusters 做 paired bootstrap，95% CI 为 [+0.000059, +0.000355]。High PCC 提升 +0.002314；High F1 基本不变。Context PCC 下降 -0.001277，因此当前选择明确偏向主权重更高的 FC，而不是宣称全指标提升。",
    )
    add_callout(
        doc,
        "如何理解这个增益",
        "数值很小，但配对区间为正，说明方向在当前 OOF 合同中稳定；它仍需要官方隐藏测试确认，不能包装成已获得的排行榜提升。",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    add_heading(doc, "4.3 负对照与失败分支", 2)
    add_table(
        doc,
        ["候选模块", "验证信号", "最终决策"],
        [
            ("菌株 SNP-MDS", "真实表示优于 shuffled：FC +0.006699；但叠加到冻结主模型后最佳权重为 0", "不进入 R01"),
            ("已见菌株 expert", "能改善较弱 general parent；叠加到 M12.0 后 α=0 最优", "拒绝叠加"),
            ("双未知联合语义", "real-real FC 0.185253，低于 zero-semantic 0.210603", "不进入 R00"),
            ("OP3 RNA 迁移", "真实表示通过独立 shuffled 对照，并改善 S1 FC", "仅进入 R10 / S1"),
        ],
        widths=[4.0, 8.7, 3.8],
        font_size=9.0,
    )
    add_body(
        doc,
        "这些结果支持一个朴素但重要的结论：外部表征本身有信息，不代表它与当前主干的残差方向兼容。模型路由应由融合后的实际增益决定，而不是由模块的新颖程度决定。",
    )

    add_heading(doc, "4.4 工程可行性与结论边界", 2)
    add_bullet(doc, "已生成 4,454 × 4,422 的绝对 log₂ 候选预测，全部数值有限；prediction SHA256 为 4179af…e8159ab。")
    add_bullet(doc, "完整路由审计覆盖 4,454 行；go-ai 204 项测试与 RNA transfer 14 项测试通过。")
    add_bullet(doc, "当前没有官方提交 ID 或官方分数；4,422 与资料中 5,243 个原始蛋白的最终列契约仍需以最新官方模板确认。")
    add_callout(
        doc,
        "上传前必须完成",
        "2026-08-15 的候选 prediction 记录为 all_released_labeled_rows refit。若最新规则规定 validation 标签只可用于选型，"
        "则必须以 split_final=train 为唯一拟合标签重新训练并生成预测；本文件不把现有候选预测视为最终合规提交。",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )

    # Section 5
    add_heading(doc, "五、复现与开放计划", 1)
    add_heading(doc, "5.1 复现方式", 2)
    add_body(doc, "复现实验以配置、manifest 和模型台账为入口，固定数据哈希、fold、随机种子、checkpoint 与输出合同。计划提供以下最小流程：")
    add_number(doc, 1, "准备环境", "安装 Python / PyTorch / RDKit 依赖，核验官方数据与 OP3 文件 SHA256。")
    add_number(doc, 2, "构建特征", "生成 sample_ID 对齐的数据合同、Morgan-2048、entity support vocabulary 与 OOF folds。")
    add_number(doc, 3, "训练与评估", "依次运行 M6、OP3 encoder / M9.6、M12 固定融合；输出逐 fold 指标、bootstrap 与负对照。")
    add_number(doc, 4, "生成预测", "按 support router 推理并运行提交校验器；保存 prediction contract、route audit 与 SHA256。")

    add_heading(doc, "5.2 开源计划", 2)
    add_body(
        doc,
        "初赛阶段提交方案文档。进入复赛后，计划公开可运行代码、环境配置、训练 / 推理入口、模型卡、实验台账与最小复现样例；官方非公开数据、原始标签、测试预测和受限缓存不进入公共仓库。若组委会要求容器化，将提供固定镜像或等价环境锁定文件。",
    )

    add_heading(doc, "5.3 依赖、数据来源与合规披露", 2)
    add_body(
        doc,
        "外部预训练仅使用 OP3 / Open Problems Single-Cell Perturbations 的公开 pseudobulk differential expression 文件（CC BY 4.0）。其来源 URL、SHA256、预训练配置和排除的 GOAI parent structures 均写入 manifest。核心建模流程未调用商业 API 或闭源模型。",
    )
    add_body(
        doc,
        "正式提交将同时披露：官方数据的使用边界、外部数据许可、软件版本、随机种子、训练标签范围、模型路由、输出列契约与预测哈希。任何验证标签边界或官方模板更新，以组委会最新规则为准并重新生成产物。",
    )

    # Section 6
    add_heading(doc, "六、团队介绍", 1)
    add_callout(
        doc,
        "待团队本人填写",
        "本节故意保留占位。请仅填写真实、可核验的信息，并统一姓名、单位与报名系统中的表述。",
        fill=PALE_GRAY,
        accent=MUTED,
    )
    add_heading(doc, "6.1 成员背景", 2)
    add_table(
        doc,
        ["成员", "学校 / 公司、岗位 / 专业", "核心技能与研究方向"],
        [
            ("成员 1【待填】", "【待填】", "【待填】"),
            ("成员 2【待填】", "【待填】", "【待填】"),
            ("成员 3【待填】", "【待填】", "【待填】"),
        ],
        widths=[3.2, 6.2, 7.1],
        font_size=9.4,
    )

    add_heading(doc, "6.2 团队分工", 2)
    add_table(
        doc,
        ["成员", "角色", "具体职责"],
        [
            ("成员 1【待填】", "【待填】", "【待填】"),
            ("成员 2【待填】", "【待填】", "【待填】"),
            ("成员 3【待填】", "【待填】", "【待填】"),
        ],
        widths=[3.2, 4.2, 9.1],
        font_size=9.4,
    )

    add_heading(doc, "6.3 团队成果", 2)
    add_table(
        doc,
        ["成员 / 团队", "项目、论文或获奖经历", "链接 / 证明材料"],
        [
            ("【待填】", "【待填】", "【待填】"),
            ("【待填】", "【待填】", "【待填】"),
        ],
        widths=[3.2, 7.7, 5.6],
        font_size=9.4,
    )

    p = doc.add_paragraph()
    set_para(p, before=18, after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "— 完 —", color=MUTED, size=9.5)

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
